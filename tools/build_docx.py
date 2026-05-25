from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DATA = ROOT / "data" / "outline.zh-CN.json"
DEFAULT_DOCX = ROOT / "dist" / "数字取证快查大纲.docx"
DEFAULT_MD = ROOT / "dist" / "数字取证快查大纲.md"


def set_run_font(run, *, name: str = "SimSun", size: float = 8.4, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.0

    heading = styles["Heading 1"]
    heading.font.name = "SimSun"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    heading.font.size = Pt(15)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(5)
    heading.paragraph_format.line_spacing = 1.0

    if "TreeText" in styles:
        tree = styles["TreeText"]
    else:
        tree = styles.add_style("TreeText", WD_STYLE_TYPE.PARAGRAPH)
    tree.font.name = "Courier New"
    tree._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    tree.font.size = Pt(8.2)
    tree.font.color.rgb = RGBColor(0, 0, 0)
    tree.paragraph_format.left_indent = Pt(12)
    tree.paragraph_format.first_line_indent = Pt(-12)
    tree.paragraph_format.space_before = Pt(0)
    tree.paragraph_format.space_after = Pt(0)
    tree.paragraph_format.line_spacing = 0.95


def remove_default_paragraph(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:p") and not "".join(child.itertext()).strip():
            body.remove(child)
            return


def page_label(page: str) -> str:
    return f"P{page}"


def add_tree_paragraph(doc: Document, prefix: str, runs: list[tuple[str, bool]]) -> None:
    para = doc.add_paragraph(style="TreeText")
    if prefix:
        run = para.add_run(prefix)
        set_run_font(run, name="Courier New", size=8.2)
    for text, bold in runs:
        run = para.add_run(text)
        set_run_font(run, name="SimSun", size=8.2, bold=bold)


def add_front_matter(doc: Document, data: dict[str, Any]) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data["title"])
    set_run_font(run, size=18, bold=True)
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(data["subtitle"])
    set_run_font(run, size=9)
    subtitle.paragraph_format.space_after = Pt(2)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(data["page_note"])
    set_run_font(run, size=9)
    note.paragraph_format.space_after = Pt(6)

    doc.add_heading("快查大纲", level=1)


def build_docx(data: dict[str, Any], output_path: Path) -> None:
    doc = Document()
    remove_default_paragraph(doc)
    configure_styles(doc)
    add_front_matter(doc, data)

    for chapter in data["chapters"]:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(5)
        heading.paragraph_format.space_after = Pt(1)
        run = heading.add_run(chapter["title"])
        set_run_font(run, size=11, bold=True)

        sections = chapter["sections"]
        for section_index, section in enumerate(sections):
            section_last = section_index == len(sections) - 1
            branch = "└─ " if section_last else "├─ "
            child_prefix = "   " if section_last else "│  "
            add_tree_paragraph(
                doc,
                branch,
                [(section["title"], True), (f" {page_label(section['page'])}", False)],
            )

            entries: list[list[tuple[str, bool]]] = []
            for item in section["items"]:
                entries.append([(item["text"], bool(item.get("emphasis"))), (f" {page_label(item['page'])}", False)])
            if section.get("terms"):
                entries.append([("核心名词", True), ("：" + section["terms"], False)])

            for item_index, runs in enumerate(entries):
                item_last = item_index == len(entries) - 1
                item_branch = "└─ " if item_last else "├─ "
                add_tree_paragraph(doc, child_prefix + item_branch, runs)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def markdown_lines(data: dict[str, Any]) -> list[str]:
    lines = [
        f"# {data['title']}",
        "",
        data["subtitle"],
        "",
        data["page_note"],
        "",
        "# 快查大纲",
        "",
    ]
    for chapter in data["chapters"]:
        lines.append(chapter["title"])
        sections = chapter["sections"]
        for section_index, section in enumerate(sections):
            section_last = section_index == len(sections) - 1
            branch = "└─" if section_last else "├─"
            child_prefix = "   " if section_last else "│  "
            lines.append(f"{branch} **{section['title']}** {page_label(section['page'])}")

            entries = [f"{item['text']} {page_label(item['page'])}" for item in section["items"]]
            if section.get("terms"):
                entries.append(f"**核心名词**：{section['terms']}")
            for item_index, entry in enumerate(entries):
                item_last = item_index == len(entries) - 1
                item_branch = "└─" if item_last else "├─"
                lines.append(f"{child_prefix}{item_branch} {entry}")
        lines.append("")
    return lines


def build_markdown(data: dict[str, Any], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(markdown_lines(data)).rstrip() + "\n", encoding="utf-8")


def load_data(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the digital forensics quick outline DOCX.")
    parser.add_argument("--data", type=Path, default=DEFAULT_DATA, help="Path to outline JSON data.")
    parser.add_argument("--output", type=Path, default=DEFAULT_DOCX, help="Output DOCX path.")
    parser.add_argument("--markdown", type=Path, default=DEFAULT_MD, help="Output Markdown path.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    data = load_data(args.data)
    build_docx(data, args.output)
    build_markdown(data, args.markdown)
    print(f"Wrote {args.output}")
    print(f"Wrote {args.markdown}")


if __name__ == "__main__":
    main()
